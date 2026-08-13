#!/usr/bin/env python3
"""Materialize the fixture corpus into its delivery formats.

The corpus is deliberately mixed-format because format is where offset-preserving parsing
actually gets tested. A corpus of fifteen markdown files would exercise one extractor and
prove nothing about the two that are hard.

Formats are assigned to essays whose content makes the format meaningful rather than at
random: the quotation-heavy paper carries an embedded chart (so `pdf_text` extraction has to
survive a figure), the statistics paper is a PDF, and the footnoted paper is a `.docx` (so the
injection payload has to survive a real extraction path before the scanner sees it).

Run after `scripts/` has the authored essay JSON. Regeneration is deterministic apart from
PDF creation timestamps, which is why the PDFs are committed rather than rebuilt in CI.
"""

from __future__ import annotations

import json
import re
import sys
from pathlib import Path

REPO = Path(__file__).resolve().parent.parent
FIXTURES = REPO / "fixtures"

# Which essay ships in which format, and why that pairing tests something.
FORMATS = {
    "s01": "md",
    "s03": "md",
    "s05": "md",
    "s08": "md",
    "s09": "md",
    "s12": "md",
    "s15": "md",
    "s02": "docx",
    "s04": "docx",
    "s07": "docx",
    "s10": "docx",
    "s13": "docx",
    "s06": "pdf",
    "s11": "pdf",
    "s14": "pdf",
}

# The three-submission iteration subset (KAR-204). Chosen for behavioural coverage rather
# than for being quick: one clean run, the injection plant, and the no-evidence plant --
# which between them exercise the accept path, the scanner, and the branch that must never
# enter the retry loop. One of each of the three formats, so the extractors stay covered too.
DEV_SUBSET = ("s01", "s07", "s12")


def markdown_to_paragraphs(body: str) -> list[str]:
    """Flatten markdown to the paragraph sequence a reader would see.

    Headings are kept as their own paragraphs. They are part of the document's visible
    structure and a citation may legitimately land on one.
    """
    blocks = [b.strip() for b in re.split(r"\n\s*\n", body) if b.strip()]
    out: list[str] = []
    for block in blocks:
        cleaned = re.sub(r"^#{1,6}\s*", "", block)
        cleaned = cleaned.replace("\n", " ")
        cleaned = re.sub(r"\s{2,}", " ", cleaned).strip()
        if cleaned and cleaned != "---":
            out.append(cleaned)
    return out


def write_md(path: Path, essay: dict) -> None:
    path.write_text(f"# {essay['title']}\n\n{essay['body'].strip()}\n", encoding="utf-8")


def write_docx(path: Path, essay: dict) -> None:
    import docx

    document = docx.Document()
    document.add_paragraph(essay["title"])
    for paragraph in markdown_to_paragraphs(essay["body"]):
        document.add_paragraph(paragraph)
    document.save(str(path))


def write_pdf(path: Path, essay: dict, *, with_chart: bool = False) -> None:
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    styles = getSampleStyleSheet()
    story: list = [Paragraph(essay["title"], styles["Title"]), Spacer(1, 0.2 * inch)]

    paragraphs = markdown_to_paragraphs(essay["body"])
    for paragraph in paragraphs:
        safe = paragraph.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        story.append(Paragraph(safe, styles["BodyText"]))
        story.append(Spacer(1, 0.10 * inch))

        # The chart goes where the essay refers to "the figure below", so the reference is
        # real rather than decorative -- an extractor that drops the figure leaves a
        # dangling reference the pipeline has to cope with.
        if with_chart and re.search(r"figure below", paragraph, re.IGNORECASE):
            story.append(_chart())
            story.append(Spacer(1, 0.15 * inch))

    SimpleDocTemplate(str(path), pagesize=LETTER, title=essay["title"]).build(story)


def _chart():
    """A small bar chart: subscription rate by household income band."""
    from reportlab.platypus import Flowable

    class ChartFlowable(Flowable):
        width, height = 400, 190

        def wrap(self, *_):
            return self.width, self.height

        def draw(self):
            c = self.canv
            bands = [("<25k", 41), ("25-50k", 58), ("50-75k", 72), ("75-100k", 84), ("100k+", 91)]
            c.setFont("Helvetica-Bold", 10)
            c.drawString(0, 175, "Figure 1. Broadband subscription rate by household income")
            c.setFont("Helvetica", 8)
            base, scale = 25, 1.35
            for i, (label, value) in enumerate(bands):
                x = 30 + i * 70
                c.setFillGray(0.45)
                c.rect(x, base, 44, value * scale, fill=1, stroke=0)
                c.setFillGray(0)
                c.drawCentredString(x + 22, base + value * scale + 4, f"{value}%")
                c.drawCentredString(x + 22, base - 11, label)
            c.line(20, base, 380, base)

    return ChartFlowable()


def write_unparseable(path: Path) -> None:
    """A file that presents as a PDF and is not one.

    Deliberately truncated mid-object with a corrupt xref, which is what a real interrupted
    upload produces. The point is not that Karani detects a fake PDF; it is that a submission
    which *cannot be read at all* becomes a visible `TaskFailed` and an anomaly-queue item,
    rather than an empty rendition that would report the student submitted nothing relevant.
    """
    path.write_bytes(
        b"%PDF-1.7\n"
        b"1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n"
        b"2 0 obj\n<< /Type /Pages /Kids [3 0 R] /Count 1 >>\nendobj\n"
        b"3 0 obj\n<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents 4 0 R >>\n"
        b"endobj\n4 0 obj\n<< /Length 512 >>\nstream\nBT /F1 12 Tf 72 720 Td (The "
        b"\xff\xfe\x00corrupted stream begins here and never term"
    )


def main() -> int:
    essays_path = Path(sys.argv[1]) if len(sys.argv) > 1 else Path("/tmp/karani_essays_v2.json")
    essays = json.loads(essays_path.read_text(encoding="utf-8"))

    FIXTURES.mkdir(parents=True, exist_ok=True)
    (FIXTURES / "dev").mkdir(exist_ok=True)

    written: list[str] = []
    for sid in sorted(essays):
        essay = essays[sid]
        fmt = FORMATS.get(sid, "md")
        path = FIXTURES / f"{sid}.{fmt}"
        if fmt == "md":
            write_md(path, essay)
        elif fmt == "docx":
            write_docx(path, essay)
        else:
            write_pdf(path, essay, with_chart=(sid == "s06"))
        written.append(path.name)

    # The unparseable submission (KAR-202). Numbered beyond the essay corpus because it has
    # no essay content: making it one of the fifteen would cost a behavioural fixture.
    write_unparseable(FIXTURES / "s16.pdf")
    written.append("s16.pdf")

    for sid in DEV_SUBSET:
        fmt = FORMATS[sid]
        source = FIXTURES / f"{sid}.{fmt}"
        (FIXTURES / "dev" / source.name).write_bytes(source.read_bytes())

    print(f"wrote {len(written)} submissions to fixtures/")
    for name in written:
        size = (FIXTURES / name).stat().st_size
        print(f"  {name:<12} {size:>8,} bytes")
    print(f"\ndev subset ({', '.join(DEV_SUBSET)}) copied to fixtures/dev/")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
